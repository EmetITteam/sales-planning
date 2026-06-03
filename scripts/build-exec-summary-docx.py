"""
Генерує docs/EXECUTIVE_SUMMARY.docx з docs/EXECUTIVE_SUMMARY.md.
Структурований Word-документ з таблицями, заголовками, списками, code-блоками.
Стиль: EMET корпоративний (синій акцент #066aab + сіра типографіка).

Usage: python scripts/build-exec-summary-docx.py
"""
import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn, nsmap
from docx.oxml import OxmlElement

EMET_BLUE = RGBColor(0x06, 0x6A, 0xAB)
EMET_MINT = RGBColor(0x5B, 0xD5, 0xBC)
SLATE_700 = RGBColor(0x33, 0x41, 0x55)
SLATE_500 = RGBColor(0x64, 0x74, 0x8B)
INK = RGBColor(0x08, 0x1E, 0x2D)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
LIGHT_BG = RGBColor(0xF1, 0xF5, 0xF9)


def set_cell_background(cell, color_hex):
    """Set table cell background color (e.g. '066aab')."""
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    tc_pr.append(shd)


def set_cell_border(cell, color_hex='cccccc'):
    """Set thin grey border on cell."""
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right'):
        b = OxmlElement(f'w:{edge}')
        b.set(qn('w:val'), 'single')
        b.set(qn('w:sz'), '4')
        b.set(qn('w:color'), color_hex)
        tc_borders.append(b)
    tc_pr.append(tc_borders)


def add_run_with_inline_format(paragraph, text, base_bold=False, base_italic=False, base_color=None, base_size=None):
    """Parse inline markdown (**bold**, *italic*, `code`) and add formatted runs to paragraph."""
    # Tokenize by ** then * then `
    pattern = r'(\*\*[^*]+\*\*|`[^`]+`|\*[^*]+\*)'
    parts = re.split(pattern, text)
    for part in parts:
        if not part:
            continue
        is_bold = base_bold
        is_italic = base_italic
        is_code = False
        content = part
        if part.startswith('**') and part.endswith('**') and len(part) > 4:
            is_bold = True
            content = part[2:-2]
        elif part.startswith('`') and part.endswith('`') and len(part) > 2:
            is_code = True
            content = part[1:-1]
        elif part.startswith('*') and part.endswith('*') and len(part) > 2:
            is_italic = True
            content = part[1:-1]
        run = paragraph.add_run(content)
        run.bold = is_bold
        run.italic = is_italic
        if is_code:
            run.font.name = 'Consolas'
            run.font.size = Pt(10)
            run.font.color.rgb = SLATE_700
        else:
            if base_size:
                run.font.size = base_size
            if base_color:
                run.font.color.rgb = base_color


def add_markdown_table(doc, header_cells, data_rows):
    table = doc.add_table(rows=1 + len(data_rows), cols=len(header_cells))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header
    for i, cell_text in enumerate(header_cells):
        cell = table.rows[0].cells[i]
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        set_cell_background(cell, '066AAB')
        set_cell_border(cell, '054D80')
        para = cell.paragraphs[0]
        para.paragraph_format.space_before = Pt(2)
        para.paragraph_format.space_after = Pt(2)
        run = para.add_run(cell_text.strip())
        run.bold = True
        run.font.color.rgb = WHITE
        run.font.size = Pt(10)

    # Data rows
    for r, row in enumerate(data_rows, start=1):
        for c, cell_text in enumerate(row):
            if c >= len(header_cells):
                continue
            cell = table.rows[r].cells[c]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_border(cell, 'D5DAE0')
            if r % 2 == 0:
                set_cell_background(cell, 'F8FAFC')
            para = cell.paragraphs[0]
            para.paragraph_format.space_before = Pt(3)
            para.paragraph_format.space_after = Pt(3)
            add_run_with_inline_format(para, cell_text.strip(), base_size=Pt(10))

    # Spacing after table
    sp = doc.add_paragraph()
    sp.paragraph_format.space_after = Pt(4)


def add_code_block(doc, lines):
    para = doc.add_paragraph()
    para.paragraph_format.left_indent = Cm(0.5)
    para.paragraph_format.space_before = Pt(4)
    para.paragraph_format.space_after = Pt(8)
    p_pr = para._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:fill'), 'F1F5F9')
    p_pr.append(shd)
    run = para.add_run('\n'.join(lines))
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    run.font.color.rgb = SLATE_700


def main():
    repo_root = Path(__file__).resolve().parent.parent
    md_path = repo_root / 'docs' / 'EXECUTIVE_SUMMARY.md'
    out_path = repo_root / 'docs' / 'EXECUTIVE_SUMMARY.docx'

    md = md_path.read_text(encoding='utf-8')

    doc = Document()

    # Page setup — A4 with comfortable margins
    section = doc.sections[0]
    section.page_height = Cm(29.7)
    section.page_width = Cm(21.0)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)

    # Default style
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)
    style.font.color.rgb = INK
    pf = style.paragraph_format
    pf.space_after = Pt(6)
    pf.line_spacing = 1.25

    lines = md.split('\n')
    i = 0
    in_code = False
    code_buf = []
    table_buf = []
    in_table = False

    def flush_table():
        nonlocal table_buf, in_table
        if not table_buf:
            in_table = False
            return
        # Parse header + separator + rows
        # table_buf[0] is header line "| a | b |"
        # table_buf[1] is separator "|---|---|"
        # table_buf[2..] are data rows
        if len(table_buf) < 2:
            in_table = False
            table_buf = []
            return
        header = [c.strip() for c in table_buf[0].strip('|').split('|')]
        rows = []
        for raw in table_buf[2:]:
            cells = [c.strip() for c in raw.strip('|').split('|')]
            rows.append(cells)
        add_markdown_table(doc, header, rows)
        table_buf = []
        in_table = False

    while i < len(lines):
        line = lines[i]
        stripped = line.rstrip()

        # Code block fence
        if stripped.startswith('```'):
            if in_code:
                add_code_block(doc, code_buf)
                code_buf = []
                in_code = False
            else:
                if in_table:
                    flush_table()
                in_code = True
            i += 1
            continue
        if in_code:
            code_buf.append(line)
            i += 1
            continue

        # Table detection
        if stripped.startswith('|') and stripped.endswith('|'):
            if in_table:
                table_buf.append(stripped)
            else:
                table_buf = [stripped]
                in_table = True
            i += 1
            continue
        elif in_table:
            flush_table()

        # Horizontal rule
        if re.match(r'^-{3,}$', stripped):
            sp = doc.add_paragraph()
            sp.paragraph_format.space_before = Pt(4)
            sp.paragraph_format.space_after = Pt(4)
            p_pr = sp._p.get_or_add_pPr()
            pb = OxmlElement('w:pBdr')
            bot = OxmlElement('w:bottom')
            bot.set(qn('w:val'), 'single')
            bot.set(qn('w:sz'), '6')
            bot.set(qn('w:color'), '066AAB')
            pb.append(bot)
            p_pr.append(pb)
            i += 1
            continue

        # Headings
        h_match = re.match(r'^(#{1,6})\s+(.*)$', stripped)
        if h_match:
            level = len(h_match.group(1))
            content = h_match.group(2)
            # Strip markdown emphasis from content
            clean = re.sub(r'\*\*([^*]+)\*\*', r'\1', content)
            clean = re.sub(r'\*([^*]+)\*', r'\1', clean)
            clean = re.sub(r'`([^`]+)`', r'\1', clean)
            if level == 1:
                # Title
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(16)
                p.paragraph_format.space_after = Pt(4)
                run = p.add_run(clean)
                run.font.size = Pt(24)
                run.font.bold = True
                run.font.color.rgb = EMET_BLUE
            elif level == 2:
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(14)
                p.paragraph_format.space_after = Pt(6)
                run = p.add_run(clean)
                run.font.size = Pt(16)
                run.font.bold = True
                run.font.color.rgb = INK
            elif level == 3:
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(10)
                p.paragraph_format.space_after = Pt(4)
                run = p.add_run(clean)
                run.font.size = Pt(13)
                run.font.bold = True
                run.font.color.rgb = SLATE_700
            else:
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(6)
                p.paragraph_format.space_after = Pt(2)
                run = p.add_run(clean)
                run.font.size = Pt(11)
                run.font.bold = True
                run.font.color.rgb = SLATE_700
            i += 1
            continue

        # Unordered list
        list_match = re.match(r'^(\s*)-\s+(.*)$', line)
        if list_match:
            indent = len(list_match.group(1))
            content = list_match.group(2)
            p = doc.add_paragraph()
            p.style = 'List Bullet'
            p.paragraph_format.left_indent = Cm(0.6 + 0.4 * (indent // 2))
            p.paragraph_format.space_after = Pt(2)
            add_run_with_inline_format(p, content)
            i += 1
            continue

        # Numbered list
        num_match = re.match(r'^(\d+)\.\s+(.*)$', line)
        if num_match:
            content = num_match.group(2)
            p = doc.add_paragraph()
            p.style = 'List Number'
            p.paragraph_format.left_indent = Cm(0.6)
            p.paragraph_format.space_after = Pt(2)
            add_run_with_inline_format(p, content)
            i += 1
            continue

        # Empty line — just spacer
        if not stripped:
            i += 1
            continue

        # Italic-only line (often _Документ оновлюється..._)
        ital_match = re.match(r'^_(.+)_$', stripped)
        if ital_match:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(8)
            run = p.add_run(ital_match.group(1))
            run.italic = True
            run.font.size = Pt(10)
            run.font.color.rgb = SLATE_500
            i += 1
            continue

        # Normal paragraph
        p = doc.add_paragraph()
        add_run_with_inline_format(p, line)
        i += 1

    if in_table:
        flush_table()
    if in_code:
        add_code_block(doc, code_buf)

    doc.save(out_path)
    print(f'OK Wrote {out_path}')
    print(f'   Size: {out_path.stat().st_size:,} bytes')


if __name__ == '__main__':
    main()
